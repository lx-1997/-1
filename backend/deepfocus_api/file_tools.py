from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import Optional

from fastapi import HTTPException, UploadFile, status

from .schemas import FileExtractionResponse


MAX_UPLOAD_BYTES = 12 * 1024 * 1024
MAX_EXTRACTED_CHARS = 80_000


async def extract_upload_file(file: UploadFile) -> FileExtractionResponse:
    raw = await file.read()
    filename = file.filename or "uploaded-file"
    suffix = Path(filename).suffix.lower()

    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="文件过大，当前限制为 12MB。",
        )

    try:
        text, parser = _extract_by_suffix(raw, suffix)
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"缺少文件解析依赖：{exc}",
        ) from exc
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"文件解析失败：{exc}",
        ) from exc

    text = _clean_text(text)
    truncated = len(text) > MAX_EXTRACTED_CHARS
    if truncated:
        text = text[:MAX_EXTRACTED_CHARS]

    return FileExtractionResponse(
        filename=filename,
        content_type=file.content_type,
        text=text,
        char_count=len(text),
        truncated=truncated,
        parser=parser,
    )


def _extract_by_suffix(raw: bytes, suffix: str) -> tuple[str, str]:
    if suffix == ".pdf":
        return _extract_pdf(raw), "pypdf"
    if suffix == ".docx":
        return _extract_docx(raw), "python-docx"
    if suffix == ".xlsx":
        return _extract_xlsx(raw), "openpyxl"
    if suffix == ".csv":
        return _extract_csv(raw), "csv"
    if suffix == ".json":
        return _extract_json(raw), "json"
    if suffix in {".txt", ".md", ".markdown", ".log"}:
        return _decode_text(raw), "text"

    # Last resort for pasted/exported plain text with an uncommon extension.
    return _decode_text(raw), "text-fallback"


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for index, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {index + 1}]\n{page_text}")
    return "\n\n".join(pages)


def _extract_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(raw))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    return "\n".join([*paragraphs, *table_rows])


def _extract_xlsx(raw: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    lines = []
    for sheet in workbook.worksheets[:8]:
        lines.append(f"[Sheet: {sheet.title}]")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
            if values:
                lines.append(" | ".join(values))
            if row_index >= 500:
                lines.append("[Sheet truncated at 500 rows]")
                break
    return "\n".join(lines)


def _extract_csv(raw: bytes) -> str:
    text = _decode_text(raw)
    reader = csv.reader(io.StringIO(text))
    lines = []
    for index, row in enumerate(reader, start=1):
        lines.append(" | ".join(cell.strip() for cell in row))
        if index >= 1000:
            lines.append("[CSV truncated at 1000 rows]")
            break
    return "\n".join(lines)


def _extract_json(raw: bytes) -> str:
    text = _decode_text(raw)
    parsed = json.loads(text)
    return json.dumps(parsed, ensure_ascii=False, indent=2)


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _clean_text(text: Optional[str]) -> str:
    if not text:
        return ""
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(lines).strip()
